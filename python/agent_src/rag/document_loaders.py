"""
文档加载器模块
支持多种文件格式：TXT、PDF、Word、Excel、CSV、Markdown 等
"""

import os
from typing import List, Optional
from langchain_core.documents import Document


class BaseLoader:
    """文档加载器基类"""

    def __init__(self, file_path: str, encoding: str = "utf-8"):
        self.file_path = file_path
        self.encoding = encoding

    def load(self) -> List[Document]:
        raise NotImplementedError


class TxtLoader(BaseLoader):
    """TXT 文件加载器"""

    def load(self) -> List[Document]:
        # 尝试多种编码
        encodings = ["utf-8", "gbk", "gb2312", "utf-16", "latin-1"]
        content = None
        used_encoding = None

        for encoding in encodings:
            try:
                with open(self.file_path, "r", encoding=encoding) as f:
                    content = f.read()
                used_encoding = encoding
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        if content is None:
            raise ValueError(f"无法读取文件，尝试了编码: {encodings}")

        # 规范化换行符
        content = content.replace('\r\n', '\n').replace('\r', '\n')

        return [Document(page_content=content, metadata={"source": self.file_path})]


class PDFLoader(BaseLoader):
    """PDF 文件加载器 - 支持多种解析方式"""

    def load(self) -> List[Document]:
        documents = []

        # 方法1: 尝试使用 pdfplumber (推荐，对中文支持好)
        try:
            import pdfplumber
            with pdfplumber.open(self.file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        documents.append(Document(
                            page_content=text,
                            metadata={"source": self.file_path, "page": i + 1, "parser": "pdfplumber"}
                        ))
            if documents:
                return documents
        except ImportError:
            pass
        except Exception as e:
            print(f"pdfplumber 解析失败: {e}")

        # 方法2: 尝试使用 PyMuPDF (fitz)
        try:
            import fitz
            doc = fitz.open(self.file_path)
            for i, page in enumerate(doc):
                text = page.get_text()
                if text and text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={"source": self.file_path, "page": i + 1, "parser": "PyMuPDF"}
                    ))
            doc.close()
            if documents:
                return documents
        except ImportError:
            pass
        except Exception as e:
            print(f"PyMuPDF 解析失败: {e}")

        # 方法3: 使用 pypdf 作为后备
        try:
            from pypdf import PdfReader
            reader = PdfReader(self.file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={"source": self.file_path, "page": i + 1, "parser": "pypdf"}
                    ))
            if documents:
                return documents
        except ImportError:
            raise ImportError("请安装 PDF 解析库: pip install pdfplumber 或 pip install PyMuPDF")
        except Exception as e:
            print(f"pypdf 解析失败: {e}")

        if not documents:
            raise ValueError("PDF 解析失败，未能提取任何文本内容。可能是扫描版 PDF，请尝试 OCR 工具。")

        return documents


class DocxLoader(BaseLoader):
    """Word 文档加载器"""

    def load(self) -> List[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = DocxDocument(self.file_path)
        documents = []

        # 提取段落
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # 合并内容
        content = "\n\n".join(paragraphs)

        if content.strip():
            documents.append(Document(
                page_content=content,
                metadata={"source": self.file_path}
            ))

        # 提取表格
        for i, table in enumerate(doc.tables):
            table_text = self._extract_table_text(table)
            if table_text.strip():
                documents.append(Document(
                    page_content=table_text,
                    metadata={
                        "source": self.file_path,
                        "type": "table",
                        "table_index": i + 1
                    }
                ))

        return documents

    def _extract_table_text(self, table) -> str:
        """提取表格文本"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)


class ExcelLoader(BaseLoader):
    """Excel 文件加载器"""

    def load(self) -> List[Document]:
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("请安装 pandas: pip install pandas openpyxl")

        documents = []

        # 读取所有工作表
        excel_file = pd.ExcelFile(self.file_path)

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)

            # 转换为文本格式
            content = f"工作表: {sheet_name}\n\n"
            content += df.to_string(index=False)

            documents.append(Document(
                page_content=content,
                metadata={
                    "source": self.file_path,
                    "sheet": sheet_name,
                    "type": "excel"
                }
            ))

        return documents


class CSVLoader(BaseLoader):
    """CSV 文件加载器"""

    def load(self) -> List[Document]:
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("请安装 pandas: pip install pandas")

        # 尝试多种编码
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        df = None

        for encoding in encodings:
            try:
                df = pd.read_csv(self.file_path, encoding=encoding)
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        if df is None:
            raise ValueError(f"无法读取 CSV 文件")

        content = df.to_string(index=False)

        return [Document(
            page_content=content,
            metadata={
                "source": self.file_path,
                "type": "csv"
            }
        )]


class MarkdownLoader(BaseLoader):
    """Markdown 文件加载器"""

    def load(self) -> List[Document]:
        # 使用 TXT 加载器读取内容
        txt_loader = TxtLoader(self.file_path, self.encoding)
        documents = txt_loader.load()

        # 添加 markdown 类型标记
        for doc in documents:
            doc.metadata["type"] = "markdown"

        return documents


class HTMLLoader(BaseLoader):
    """HTML 文件加载器"""

    def load(self) -> List[Document]:
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("请安装 beautifulsoup4: pip install beautifulsoup4")

        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        content = None

        for encoding in encodings:
            try:
                with open(self.file_path, "r", encoding=encoding) as f:
                    html_content = f.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        soup = BeautifulSoup(html_content, 'html.parser')

        # 移除脚本和样式
        for script in soup(["script", "style"]):
            script.decompose()

        # 获取文本
        text = soup.get_text(separator='\n')
        # 清理多余空白
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        content = "\n".join(lines)

        return [Document(
            page_content=content,
            metadata={
                "source": self.file_path,
                "type": "html"
            }
        )]


class JSONLoader(BaseLoader):
    """JSON 文件加载器"""

    def load(self) -> List[Document]:
        import json

        with open(self.file_path, "r", encoding=self.encoding) as f:
            data = json.load(f)

        # 将 JSON 转换为格式化文本
        content = json.dumps(data, ensure_ascii=False, indent=2)

        return [Document(
            page_content=content,
            metadata={
                "source": self.file_path,
                "type": "json"
            }
        )]


# 文件扩展名到加载器的映射
LOADER_MAPPING = {
    ".txt": TxtLoader,
    ".pdf": PDFLoader,
    ".docx": DocxLoader,
    ".doc": DocxLoader,  # 注意: .doc 格式可能需要额外处理
    ".xlsx": ExcelLoader,
    ".xls": ExcelLoader,
    ".csv": CSVLoader,
    ".md": MarkdownLoader,
    ".markdown": MarkdownLoader,
    ".html": HTMLLoader,
    ".htm": HTMLLoader,
    ".json": JSONLoader,
}


def get_loader(file_path: str, encoding: str = "utf-8") -> BaseLoader:
    """
    根据文件扩展名获取对应的加载器

    Args:
        file_path: 文件路径
        encoding: 文件编码

    Returns:
        对应的文档加载器实例

    Raises:
        ValueError: 不支持的文件格式
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext not in LOADER_MAPPING:
        raise ValueError(f"不支持的文件格式: {ext}。支持的格式: {list(LOADER_MAPPING.keys())}")

    loader_class = LOADER_MAPPING[ext]
    return loader_class(file_path, encoding)


def load_document(file_path: str, encoding: str = "utf-8") -> List[Document]:
    """
    加载文档（自动识别格式）

    Args:
        file_path: 文件路径
        encoding: 文件编码

    Returns:
        Document 对象列表
    """
    loader = get_loader(file_path, encoding)
    return loader.load()


def get_supported_formats() -> List[str]:
    """获取支持的文件格式列表"""
    return list(LOADER_MAPPING.keys())
